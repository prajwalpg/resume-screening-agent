"""DOCX text extraction built on python-docx (paragraphs + tables)."""

from pathlib import Path
from typing import Union


def extract_docx_text(path: Union[str, Path]) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "python-docx is required to parse DOCX files. Install it with: pip install python-docx"
        ) from exc

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    # DOCX resumes sometimes keep skills/education inside tables.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts).strip()
