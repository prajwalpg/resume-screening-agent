"""Generate the sample evaluation dataset: 1 job description + 12 resumes.

The resumes are written across all three supported formats (PDF / DOCX / TXT)
and deliberately span the full scoring spectrum -- strong matches, partial
matches, weak matches, a fresher with no experience, and one exact duplicate
(to exercise duplicate detection).

Run from the project root:
    python scripts/generate_sample_data.py
"""

import sys
from pathlib import Path
from typing import Dict, List

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

RESUMES_DIR = PROJECT_ROOT / "data" / "resumes"
JD_DIR = PROJECT_ROOT / "data" / "jd"
JD_PATH = JD_DIR / "software_test_engineer.txt"

JD_TEXT = """Job Title: Software Test Automation Engineer

About the Role
We are looking for a Software Test Automation Engineer to design, build and
maintain automated test suites for our web applications and REST APIs. You
will work with the development team in an agile environment to ensure product
quality across every release.

Responsibilities:
- Design, develop and maintain automated test scripts using Python and Selenium
- Build and run API test suites for REST services using Postman and Python
- Integrate automated tests into CI/CD pipelines (Jenkins)
- Perform regression, smoke and integration testing before every release
- Write and maintain test plans, test cases and defect reports
- Collaborate with developers to reproduce and resolve defects
- Track quality metrics and report test coverage to stakeholders

Required Skills:
- Python (strong scripting skills)
- Selenium WebDriver for web UI automation
- API Testing (REST APIs, Postman or similar tools)
- SQL for database validation
- Git for version control

Preferred Skills:
- Pytest
- Jenkins
- Docker

Minimum Experience:
1+ year of hands-on test automation experience

Education:
Bachelor's degree in Computer Science, Information Technology or a related field.
"""

# ---------------------------------------------------------------------------
# Resume definitions. `format` chooses the output file type.
# Candidate 11 is an exact duplicate of candidate 02 on purpose.
# ---------------------------------------------------------------------------
RESUMES: List[Dict] = [
    {
        "file": "candidate_01",
        "format": "pdf",
        "name": "PRIYA SHARMA",
        "email": "priya.sharma@example.com",
        "phone": "+91 98765 43210",
        "location": "Bengaluru, India",
        "summary": "QA Automation Engineer with 3 years of experience in test automation, API testing and CI/CD-driven quality assurance for web applications.",
        "skills": {
            "Programming": ["Python", "SQL", "Bash"],
            "Test Automation": ["Selenium", "Pytest", "API Testing", "Postman", "Regression Testing"],
            "CI/CD & Tools": ["Git", "Jenkins", "Docker", "Jira", "Agile"],
        },
        "experience": [
            {
                "title": "QA Automation Engineer",
                "company": "TechNova Solutions",
                "period": "2022 - Present",
                "bullets": [
                    "Built and maintained Selenium-Pytest automation suites covering 500+ regression scenarios",
                    "Automated REST API tests with Python and Postman, integrated into Jenkins CI pipelines",
                    "Containerized test environments with Docker to keep builds reproducible",
                ],
            },
            {
                "title": "QA Engineer",
                "company": "BrightApps Pvt Ltd",
                "period": "2021 - 2022",
                "bullets": [
                    "Executed manual and automated regression cycles for e-commerce releases",
                    "Logged and tracked defects in Jira and verified fixes with developers",
                ],
            },
        ],
        "projects": [
            "Hybrid Test Automation Framework: Selenium + Pytest with Allure reporting and parallel execution",
            "API Regression Suite: 300+ endpoints automated in Python with Jenkins CI integration",
        ],
        "education": ["B.Tech in Computer Science, Visvesvaraya Technological University (2017 - 2021)"],
        "certifications": ["ISTQB Certified Tester Foundation Level", "Postman API Testing Certification"],
    },
    {
        "file": "candidate_02",
        "format": "pdf",
        "name": "RAHUL KUMAR",
        "email": "rahul.kumar@example.com",
        "phone": "+91 98200 12345",
        "location": "Pune, India",
        "summary": "Software Test Engineer with 2.5 years of experience in Python-based web and API automation, focused on shifting quality left.",
        "skills": {
            "Programming": ["Python", "SQL"],
            "Test Automation": ["Selenium", "Pytest", "API Testing", "Rest Assured"],
            "Tools": ["Git", "Jenkins", "Jira", "Agile"],
        },
        "experience": [
            {
                "title": "Software Test Engineer",
                "company": "InfoEdge Systems",
                "period": "2023 - Present",
                "bullets": [
                    "Developed Selenium-Pytest automation frameworks for three client web applications",
                    "Automated API regression using Rest Assured and Python requests",
                    "Set up nightly Jenkins jobs with failure triage reports",
                ],
            },
            {
                "title": "SDET Intern",
                "company": "CloudNine Labs",
                "period": "2022 - 2023",
                "bullets": [
                    "Wrote pytest unit and integration tests for internal tools",
                    "Assisted in migrating manual test cases into automated suites",
                ],
            },
        ],
        "projects": [
            "Page Object Model framework with Selenium and Pytest (reduced maintenance effort by 40%)",
            "SQL data validation scripts for ETL pipeline checks",
        ],
        "education": ["B.E. in Computer Science, Savitribai Phule Pune University (2019 - 2023)"],
        "certifications": ["ISTQB Certified Tester Foundation Level"],
    },
    {
        "file": "candidate_03",
        "format": "docx",
        "name": "ANANYA RAO",
        "email": "ananya.rao@example.com",
        "phone": "+91 90000 23456",
        "location": "Hyderabad, India",
        "summary": "QA Engineer with around 2 years of experience in manual and web automation testing, looking to deepen her automation skill set.",
        "skills": {
            "Programming": ["Python", "SQL"],
            "Testing": ["Selenium", "Manual Testing", "Regression Testing", "Test Cases"],
            "Tools": ["Git", "Docker", "Jira", "Agile"],
        },
        "experience": [
            {
                "title": "QA Engineer",
                "company": "Zenith Retail",
                "period": "2024 - Present",
                "bullets": [
                    "Automated checkout-flow regression tests with Selenium and Python",
                    "Wrote detailed test cases and defect reports for release sign-off",
                ],
            },
            {
                "title": "Junior QA Analyst",
                "company": "StartupHub Technologies",
                "period": "2023 - 2024",
                "bullets": [
                    "Performed manual exploratory and smoke testing",
                    "Maintained the bug triage board in Jira",
                ],
            },
        ],
        "projects": ["Selenium smoke suite covering 30 critical storefront journeys"],
        "education": ["B.Tech in Information Technology, JNTU Hyderabad (2019 - 2023)"],
        "certifications": ["Coursera Test Automation Certificate"],
    },
    {
        "file": "candidate_04",
        "format": "pdf",
        "name": "KIRAN PATEL",
        "email": "kiran.patel@example.com",
        "phone": "+91 98111 34567",
        "location": "Ahmedabad, India",
        "summary": "Test Automation Engineer with 2 years of experience in Java-based web automation and database validation.",
        "skills": {
            "Programming": ["Java", "SQL"],
            "Test Automation": ["Selenium", "TestNG", "Maven"],
            "Tools": ["Git", "Jira", "Agile", "Bug Tracking"],
        },
        "experience": [
            {
                "title": "QA Automation Engineer",
                "company": "WebWorks Solutions",
                "period": "2023 - Present",
                "bullets": [
                    "Maintained Selenium-TestNG suites for a banking portal",
                    "Performed SQL validation of backend data for payment flows",
                ],
            },
            {
                "title": "QA Trainee",
                "company": "WebWorks Solutions",
                "period": "2022 - 2023",
                "bullets": ["Executed regression packs and reported defects"],
            },
        ],
        "projects": ["Data-driven TestNG framework with Excel input for 200+ scenarios"],
        "education": ["B.E. in Information Technology, Gujarat Technological University (2018 - 2022)"],
        "certifications": [],
    },
    {
        "file": "candidate_05",
        "format": "docx",
        "name": "ARJUN SINGH",
        "email": "arjun.singh@example.com",
        "phone": "+91 99887 45678",
        "location": "Noida, India",
        "summary": "Backend developer with 1.5 years of experience building Python web services, exploring a transition into test automation.",
        "skills": {
            "Programming": ["Python", "JavaScript"],
            "Frameworks": ["Django", "Flask"],
            "Cloud & Tools": ["AWS", "Docker", "Git"],
        },
        "experience": [
            {
                "title": "Software Engineer",
                "company": "Finlytics",
                "period": "2024 - Present",
                "bullets": [
                    "Developed Django services for reporting dashboards",
                    "Wrote unit tests for business logic modules",
                ],
            },
            {
                "title": "SDE Intern",
                "company": "Finlytics",
                "period": "2023 - 2024",
                "bullets": ["Fixed bugs and delivered small features"],
            },
        ],
        "projects": ["Reporting dashboard service built with Django and AWS"],
        "education": ["B.Tech in Computer Science, Dr. A.P.J. Abdul Kalam Technical University (2019 - 2023)"],
        "certifications": [],
    },
    {
        "file": "candidate_06",
        "format": "pdf",
        "name": "MEERA IYER",
        "email": "meera.iyer@example.com",
        "phone": "+91 97654 56789",
        "location": "Chennai, India",
        "summary": "Manual QA analyst with 2 years of experience in functional, regression and API testing, transitioning towards automation.",
        "skills": {
            "Testing": ["Manual Testing", "Test Cases", "Bug Tracking", "Regression Testing", "API Testing", "Postman"],
            "Database & Tools": ["SQL", "Jira", "Agile", "Git"],
        },
        "experience": [
            {
                "title": "QA Analyst",
                "company": "RetailSoft",
                "period": "2023 - Present",
                "bullets": [
                    "Designed and executed 800+ functional test cases",
                    "Tested REST APIs with Postman and validated response schemas",
                    "Verified database records with SQL queries",
                ],
            },
            {
                "title": "QA Intern",
                "company": "RetailSoft",
                "period": "2022 - 2023",
                "bullets": ["Supported UAT cycles and bug verification"],
            },
        ],
        "projects": ["Postman collection library covering 150+ endpoints"],
        "education": ["B.Sc in Computer Science, University of Madras (2019 - 2022)"],
        "certifications": ["ISTQB Certified Tester Foundation Level"],
    },
    {
        "file": "candidate_07",
        "format": "txt",
        "name": "VIKRAM JOSHI",
        "email": "vikram.joshi@example.com",
        "phone": "+91 91234 67890",
        "location": "Mumbai, India",
        "summary": "Frontend developer with 1 year of experience building responsive web applications in React.",
        "skills": {
            "Languages": ["JavaScript", "HTML", "CSS"],
            "Frameworks": ["React", "Node.js", "Express.js"],
            "Databases": ["MongoDB"],
            "Tools": ["Git", "Figma", "VS Code"],
        },
        "experience": [
            {
                "title": "Frontend Developer",
                "company": "PixelCraft Studios",
                "period": "2024 - Present",
                "bullets": [
                    "Built React components for client dashboards",
                    "Integrated REST endpoints and improved page load times",
                ],
            },
        ],
        "projects": ["Personal portfolio site with React and CSS animations"],
        "education": ["BCA in Computer Applications, Mumbai University (2021 - 2024)"],
        "certifications": [],
    },
    {
        "file": "candidate_08",
        "format": "pdf",
        "name": "SNEHA REDDY",
        "email": "sneha.reddy@example.com",
        "phone": "+91 96543 78901",
        "location": "Bengaluru, India",
        "summary": "API Automation Engineer with 2 years of experience testing REST services with Python, Pytest and Postman.",
        "skills": {
            "Programming": ["Python", "SQL"],
            "Testing": ["API Testing", "Pytest", "Postman", "Rest Assured", "Unit Testing"],
            "Tools": ["Git", "Jenkins", "Jira"],
        },
        "experience": [
            {
                "title": "API Test Engineer",
                "company": "PayFast Technologies",
                "period": "2023 - Present",
                "bullets": [
                    "Automated payment REST API regression with Pytest",
                    "Validated asynchronous workflows and database state with SQL",
                    "Integrated API suites into Jenkins pipelines",
                ],
            },
            {
                "title": "QA Intern",
                "company": "PayFast Technologies",
                "period": "2022 - 2023",
                "bullets": ["Tested billing APIs manually with Postman"],
            },
        ],
        "projects": ["Pytest-based API framework with schema validation and HTML reports"],
        "education": ["B.Tech in Computer Science, Osmania University (2019 - 2023)"],
        "certifications": ["Postman Student Expert Certification"],
    },
    {
        "file": "candidate_09",
        "format": "docx",
        "name": "ROHIT VERMA",
        "email": "rohit.verma@example.com",
        "phone": "+91 93456 89012",
        "location": "Kolkata, India",
        "summary": "Recent B.Sc Statistics graduate with a 3-month data analytics internship, eager to start a career in technology.",
        "skills": {
            "Analytics": ["Python", "Pandas", "Excel", "Data Analysis"],
            "Tools": ["Power BI"],
        },
        "experience": [
            {
                "title": "Data Analyst Intern",
                "company": "MarketPulse",
                "period": "2024 - 2024",
                "bullets": [
                    "Cleaned survey data with Pandas",
                    "Built Power BI dashboards for the marketing team",
                ],
            },
        ],
        "projects": ["Sales trend analysis notebook (Pandas, Matplotlib)"],
        "education": ["B.Sc in Statistics, Presidency University (2021 - 2024)"],
        "certifications": [],
    },
    {
        "file": "candidate_10",
        "format": "pdf",
        "name": "DIVYA NAIR",
        "email": "divya.nair@example.com",
        "phone": "+91 98760 89023",
        "location": "Kochi, India",
        "summary": "SDET with over 2.5 years of experience in UI and API automation, BDD and CI/CD quality gates.",
        "skills": {
            "Programming": ["Python", "JavaScript"],
            "Automation": ["Selenium", "Pytest", "API Testing", "Postman", "Cucumber"],
            "DevOps & Tools": ["Git", "Docker", "Jenkins", "Jira", "Agile"],
        },
        "experience": [
            {
                "title": "SDET",
                "company": "CloudLeap Technologies",
                "period": "2023 - Present",
                "bullets": [
                    "Owns the Selenium-Pytest UI framework for a SaaS product",
                    "Automated REST API suites and wired them into Jenkins",
                    "Runs Dockerized test environments for nightly builds",
                ],
            },
            {
                "title": "QA Engineer",
                "company": "CodeOrbit",
                "period": "2022 - 2023",
                "bullets": ["Automated regression for booking workflows"],
            },
        ],
        "projects": [
            "BDD Cucumber suite for checkout journeys",
            "Dockerized Selenium Grid setup for parallel runs",
        ],
        "education": ["B.Tech in Computer Science, Cochin University of Science and Technology (2018 - 2022)"],
        "certifications": ["ISTQB Agile Tester Extension"],
    },
    {
        "file": "candidate_11",  # Exact duplicate of candidate_02 (tests duplicate detection)
        "format": "pdf",
        "copy_of": "candidate_02",
    },
    {
        "file": "candidate_12",
        "format": "txt",
        "name": "KARTHIK MENON",
        "email": "karthik.menon@example.com",
        "phone": "+91 90909 90123",
        "location": "Coimbatore, India",
        "summary": "QA engineer with 2 years of experience in Java-Selenium automation for enterprise applications.",
        "skills": {
            "Programming": ["Java", "SQL"],
            "Automation": ["Selenium", "TestNG", "JUnit", "Maven"],
            "Tools": ["Git", "Jenkins", "Jira", "Agile"],
        },
        "experience": [
            {
                "title": "QA Engineer",
                "company": "ERPSoft",
                "period": "2023 - Present",
                "bullets": [
                    "Maintains Selenium-TestNG regression suites for an ERP product",
                    "Validates backend data with SQL scripts",
                ],
            },
            {
                "title": "QA Trainee",
                "company": "ERPSoft",
                "period": "2022 - 2023",
                "bullets": ["Executed smoke and regression cycles"],
            },
        ],
        "projects": ["TestNG suite refactor with retry analyzer and extent reports"],
        "education": ["B.Tech in Electronics and Communication Engineering, Anna University (2018 - 2022)"],
        "certifications": [],
    },
]


# ---------------------------------------------------------------------------
# Rendering helpers
# ---------------------------------------------------------------------------
def resume_lines(resume: Dict) -> List[str]:
    lines = [resume["name"]]
    contact = f"Email: {resume['email']} | Phone: {resume['phone']}"
    lines.append(contact)
    lines.append(f"Location: {resume['location']}")
    lines.append("")
    lines.append("PROFESSIONAL SUMMARY")
    lines.append(resume["summary"])
    lines.append("")
    lines.append("TECHNICAL SKILLS")
    for group, items in resume["skills"].items():
        lines.append(f"- {group}: {', '.join(items)}")
    lines.append("")
    lines.append("WORK EXPERIENCE")
    for role in resume["experience"]:
        lines.append(f"{role['title']}, {role['company']} ({role['period']})")
        for bullet in role["bullets"]:
            lines.append(f"- {bullet}")
        lines.append("")
    lines.append("PROJECTS")
    for project in resume["projects"]:
        lines.append(f"- {project}")
    lines.append("")
    lines.append("EDUCATION")
    for entry in resume["education"]:
        lines.append(f"- {entry}")
    lines.append("")
    lines.append("CERTIFICATIONS")
    for cert in resume.get("certifications", []) or ["None"]:
        lines.append(f"- {cert}")
    return lines


def _wrap(text: str, width: int = 90) -> List[str]:
    words = text.split()
    if not words:
        return [""]
    chunks: List[str] = []
    current = words[0]
    for word in words[1:]:
        if len(current) + 1 + len(word) <= width:
            current += " " + word
        else:
            chunks.append(current)
            current = word
    chunks.append(current)
    return chunks


def write_pdf(path: Path, lines: List[str]) -> None:
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    margin = 57.0
    y = 72.0
    bottom = 795.0
    first = True

    for raw in lines:
        is_heading = raw.isupper() and len(raw.strip()) > 3
        for chunk in _wrap(raw):
            if y > bottom:
                page = doc.new_page()
                y = 72.0
            size = 13.5 if first else 10.5
            page.insert_text((margin, y), chunk, fontsize=size, fontname="helv")
            y += 17.0 if first else 15.0
            first = False
        if is_heading:
            y += 5
        elif raw == "":
            y += 6

    doc.save(str(path))
    doc.close()


def write_docx(path: Path, lines: List[str]) -> None:
    from docx import Document

    doc = Document()
    for index, line in enumerate(lines):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        if index == 0 or (line.isupper() and len(line.strip()) > 3):
            run.bold = True
    doc.save(str(path))


def write_txt(path: Path, lines: List[str]) -> None:
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    JD_DIR.mkdir(parents=True, exist_ok=True)

    # Job description
    JD_PATH.write_text(JD_TEXT, encoding="utf-8")
    print(f"  ✓ {JD_PATH.relative_to(PROJECT_ROOT)}")

    by_key = {resume["file"]: resume for resume in RESUMES}
    writers = {"pdf": write_pdf, "docx": write_docx, "txt": write_txt}

    for entry in RESUMES:
        if "copy_of" in entry:
            # Duplicate content (same candidate), different file name.
            resume = dict(by_key[entry["copy_of"]])
            resume["file"] = entry["file"]
            resume["format"] = entry["format"]
        else:
            resume = entry
        lines = resume_lines(resume)
        target = RESUMES_DIR / f"{resume['file']}.{resume['format']}"
        writers[resume["format"]](target, lines)
        print(f"  ✓ {target.relative_to(PROJECT_ROOT)} ({target.stat().st_size} bytes)")

    print(f"\nDone: {JD_PATH.name} + {len(RESUMES)} resumes in {RESUMES_DIR}")


if __name__ == "__main__":
    main()
